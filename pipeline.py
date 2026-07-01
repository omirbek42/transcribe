"""
Пайплайн транскрибации аудио/видео (RU/KZ/EN):

  файл -> (ffmpeg) извлечение аудио -> (faster-whisper) ASR
        -> (pyannote.audio, опционально) диаризация спикеров
        -> сборка сырого транскрипта с таймкодами
        -> (Anthropic API) LLM-коррекция под деловой стиль и глоссарий
        -> экспорт .docx / .txt / .srt

CLI:
    python pipeline.py meeting.mp4 --output-dir out --diarize
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import os
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Optional

AUDIO_EXTS = {".mp3", ".wav", ".m4a"}
VIDEO_EXTS = {".mp4"}

# Порог низкой уверенности ASR-сегмента: помечаем как потенциально
# неразборчивый вместо того, чтобы додумывать текст (см. F9 в PRD).
LOW_CONFIDENCE_LOGPROB = -1.0
HIGH_NO_SPEECH_PROB = 0.6

DEFAULT_LLM_MODEL = "claude-sonnet-4-6"

SYSTEM_PROMPT = """Ты — редактор-корректор транскриптов для корпоративного документооборота
АО «Самрук-Қазына» и группы компаний фонда. Тебе передают сырой автоматический
транскрипт совещания/интервью на русском, казахском или смешанном
русско-казахском языке (переключение кодов внутри одного предложения — обычная
практика, не ошибка).

Твоя задача:
1. Исправь ошибки распознавания речи, опираясь на контекст (омофоны,
   неправильно услышанные термины, слитые/разорванные слова).
2. Приведи текст к официально-деловому стилю: убери слова-паразиты, повторы,
   незаконченные фразы — но НЕ меняй смысл и не добавляй информацию, которой
   не было в оригинале.
3. Корректно оформляй названия и термины по глоссарию ниже, расширяй по
   аналогии, если встречаешь похожие названия ДЗО фонда.
4. Названия госорганов РК пиши полностью при первом упоминании, далее —
   принятой аббревиатурой.
5. Сохраняй структуру: если это протокол совещания — оформи по схеме
   «Спикер — реплика», сохрани таймкоды в квадратных скобках в начале
   каждого блока.
6. Если фрагмент нечитаем/неразборчив — помечай [неразборчиво], не выдумывай
   содержание.
7. Не переводи казахские фрагменты на русский и наоборот — сохраняй язык
   оригинала реплики.
8. Выведи только исправленный текст, без комментариев и пояснений.

Глоссарий терминов (JSON, следуй строго):
{glossary_json}"""

ProgressCallback = Optional[Callable[[str], None]]


@dataclass
class Segment:
    start: float
    end: float
    text: str
    low_confidence: bool = False
    speaker: Optional[str] = None


def _log(progress_cb: ProgressCallback, message: str) -> None:
    if progress_cb:
        progress_cb(message)


# --------------------------------------------------------------------------
# Извлечение аудио
# --------------------------------------------------------------------------

def extract_audio(input_path: Path, workdir: Path) -> Path:
    """Извлекает/нормализует аудиодорожку в 16kHz mono WAV через ffmpeg."""
    if shutil.which("ffmpeg") is None:
        raise RuntimeError(
            "ffmpeg не найден в PATH. Установите ffmpeg (см. README) и повторите."
        )
    suffix = input_path.suffix.lower()
    if suffix not in AUDIO_EXTS | VIDEO_EXTS:
        raise ValueError(f"Неподдерживаемый формат файла: {suffix}")

    workdir.mkdir(parents=True, exist_ok=True)
    out_path = workdir / "audio_16k_mono.wav"
    cmd = [
        "ffmpeg", "-y", "-i", str(input_path),
        "-ac", "1", "-ar", "16000", "-vn",
        str(out_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"ffmpeg завершился с ошибкой:\n{result.stderr}")
    return out_path


# --------------------------------------------------------------------------
# ASR (faster-whisper)
# --------------------------------------------------------------------------

def detect_device_and_model() -> tuple[str, str, str]:
    """Возвращает (device, model_size, compute_type) по доступности CUDA."""
    try:
        import torch
        cuda_available = torch.cuda.is_available()
    except ImportError:
        cuda_available = False

    if cuda_available:
        return "cuda", "large-v3", "float16"
    return "cpu", "medium", "int8"


def transcribe(
    audio_path: Path,
    device: str,
    model_size: str,
    compute_type: str,
    progress_cb: ProgressCallback = None,
) -> list[Segment]:
    from faster_whisper import WhisperModel

    _log(progress_cb, f"Загрузка ASR-модели {model_size} ({device})...")
    model = WhisperModel(model_size, device=device, compute_type=compute_type)

    _log(progress_cb, "Распознавание речи...")
    # language=None + condition_on_previous_text=False: язык переопределяется
    # для каждого внутреннего окна декодирования, а не фиксируется на весь
    # файл — это ближе к требованию F3/F4 (переключение кодов ru/kk внутри
    # записи) при MVP-ограничениях faster-whisper.
    segments_iter, _info = model.transcribe(
        str(audio_path),
        language=None,
        vad_filter=True,
        condition_on_previous_text=False,
    )

    segments: list[Segment] = []
    for seg in segments_iter:
        low_conf = (
            seg.avg_logprob < LOW_CONFIDENCE_LOGPROB
            or seg.no_speech_prob > HIGH_NO_SPEECH_PROB
        )
        text = seg.text.strip()
        if low_conf and not text:
            text = "[неразборчиво]"
        segments.append(Segment(start=seg.start, end=seg.end, text=text, low_confidence=low_conf))
    return segments


# --------------------------------------------------------------------------
# Диаризация (pyannote.audio)
# --------------------------------------------------------------------------

def diarize(audio_path: Path, hf_token: str, progress_cb: ProgressCallback = None) -> list[tuple[float, float, str]]:
    from pyannote.audio import Pipeline

    _log(progress_cb, "Загрузка модели диаризации...")
    pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-3.1", use_auth_token=hf_token
    )

    _log(progress_cb, "Диаризация спикеров...")
    diarization = pipeline(str(audio_path))

    turns = []
    for turn, _, speaker in diarization.itertracks(yield_label=True):
        turns.append((turn.start, turn.end, speaker))
    return turns


def assign_speakers(segments: list[Segment], turns: list[tuple[float, float, str]]) -> list[Segment]:
    if not turns:
        return segments
    for seg in segments:
        mid = (seg.start + seg.end) / 2
        best_speaker = None
        best_overlap = 0.0
        for t_start, t_end, speaker in turns:
            overlap = min(seg.end, t_end) - max(seg.start, t_start)
            if overlap > best_overlap:
                best_overlap = overlap
                best_speaker = speaker
        seg.speaker = best_speaker or "Спикер"
    return segments


# --------------------------------------------------------------------------
# Сборка сырого транскрипта
# --------------------------------------------------------------------------

def format_timecode(seconds: float) -> str:
    total = int(max(0, seconds))
    hh, rem = divmod(total, 3600)
    mm, ss = divmod(rem, 60)
    return f"{hh:02d}:{mm:02d}:{ss:02d}"


def build_raw_transcript(segments: list[Segment]) -> str:
    lines = []
    for seg in segments:
        tc = format_timecode(seg.start)
        speaker = seg.speaker or "Спикер"
        lines.append(f"[{tc}] {speaker}: {seg.text}")
    return "\n".join(lines)


# --------------------------------------------------------------------------
# Глоссарий
# --------------------------------------------------------------------------

def load_glossary(path: Path) -> dict:
    if not path.exists():
        return {"companies": [], "abbreviations": {}, "people": []}
    with open(path, encoding="utf-8") as f:
        return json.load(f)


# --------------------------------------------------------------------------
# LLM-коррекция (Anthropic API)
# --------------------------------------------------------------------------

def correct_transcript(
    raw_transcript: str,
    glossary: dict,
    api_key: str,
    model: str = DEFAULT_LLM_MODEL,
    progress_cb: ProgressCallback = None,
) -> str:
    from anthropic import Anthropic

    _log(progress_cb, "LLM-коррекция транскрипта...")
    client = Anthropic(api_key=api_key)
    system = SYSTEM_PROMPT.format(
        glossary_json=json.dumps(glossary, ensure_ascii=False, indent=2)
    )

    response = client.messages.create(
        model=model,
        max_tokens=8192,
        system=system,
        messages=[{"role": "user", "content": f"Сырой транскрипт:\n\n{raw_transcript}"}],
    )
    return "".join(
        block.text for block in response.content if block.type == "text"
    ).strip()


# --------------------------------------------------------------------------
# Разбор скорректированного текста на блоки для экспорта
# --------------------------------------------------------------------------

_BLOCK_RE = re.compile(r"^\[(\d{2}):(\d{2}):(\d{2})\]\s*(?:([^:\n]+):\s*)?(.*)$")


def parse_blocks(corrected_text: str) -> list[dict]:
    blocks: list[dict] = []
    current: Optional[dict] = None
    for raw_line in corrected_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        m = _BLOCK_RE.match(line)
        if m:
            if current:
                blocks.append(current)
            hh, mm, ss, speaker, text = m.groups()
            start = int(hh) * 3600 + int(mm) * 60 + int(ss)
            current = {"start": start, "speaker": (speaker or "").strip(), "text": text.strip()}
        elif current is not None:
            current["text"] = (current["text"] + " " + line).strip()
        else:
            blocks.append({"start": None, "speaker": "", "text": line})
    if current:
        blocks.append(current)
    return blocks


# --------------------------------------------------------------------------
# Экспорт
# --------------------------------------------------------------------------

def export_txt(corrected_text: str, out_path: Path) -> None:
    out_path.write_text(corrected_text, encoding="utf-8")


def export_docx(blocks: list[dict], out_path: Path, title: str = "Протокол совещания") -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Дата формирования: {dt.date.today().isoformat()}")
    doc.add_paragraph("")

    for b in blocks:
        p = doc.add_paragraph()
        if b["start"] is not None:
            run = p.add_run(f"[{format_timecode(b['start'])}] ")
            run.bold = True
        if b["speaker"]:
            run2 = p.add_run(f"{b['speaker']}: ")
            run2.bold = True
        p.add_run(b["text"])

    doc.save(out_path)


def _srt_timecode(seconds: float) -> str:
    total_ms = int(round(seconds * 1000))
    hh, rem = divmod(total_ms, 3600_000)
    mm, rem = divmod(rem, 60_000)
    ss, ms = divmod(rem, 1000)
    return f"{hh:02d}:{mm:02d}:{ss:02d},{ms:03d}"


def export_srt(blocks: list[dict], out_path: Path) -> None:
    timed = [b for b in blocks if b["start"] is not None]
    lines = []
    for i, b in enumerate(timed):
        start = b["start"]
        end = timed[i + 1]["start"] if i + 1 < len(timed) else start + 5
        if end <= start:
            end = start + 1
        text = f"{b['speaker']}: {b['text']}" if b["speaker"] else b["text"]
        lines.append(str(i + 1))
        lines.append(f"{_srt_timecode(start)} --> {_srt_timecode(end)}")
        lines.append(text)
        lines.append("")
    out_path.write_text("\n".join(lines), encoding="utf-8")


# --------------------------------------------------------------------------
# Оркестрация пайплайна
# --------------------------------------------------------------------------

def run_pipeline(
    input_path: Path,
    output_dir: Path,
    diarize_flag: bool = False,
    glossary_path: Path = Path("glossary.json"),
    skip_llm: bool = False,
    progress_cb: ProgressCallback = None,
) -> dict[str, Path]:
    input_path = Path(input_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    workdir = output_dir / "_work"

    _log(progress_cb, "Извлечение аудио...")
    audio_path = extract_audio(input_path, workdir)

    device, model_size, compute_type = detect_device_and_model()
    segments = transcribe(audio_path, device, model_size, compute_type, progress_cb)

    if diarize_flag:
        hf_token = os.environ.get("HF_TOKEN")
        if not hf_token:
            raise RuntimeError("Диаризация требует переменную окружения HF_TOKEN.")
        turns = diarize(audio_path, hf_token, progress_cb)
        segments = assign_speakers(segments, turns)

    raw_transcript = build_raw_transcript(segments)
    glossary = load_glossary(glossary_path)

    if skip_llm:
        corrected_text = raw_transcript
    else:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise RuntimeError("LLM-коррекция требует переменную окружения ANTHROPIC_API_KEY.")
        corrected_text = correct_transcript(raw_transcript, glossary, api_key, progress_cb=progress_cb)

    blocks = parse_blocks(corrected_text)

    _log(progress_cb, "Экспорт результатов...")
    docx_path = output_dir / f"{input_path.stem}_protocol.docx"
    txt_path = output_dir / f"{input_path.stem}_transcript.txt"
    srt_path = output_dir / f"{input_path.stem}_subtitles.srt"

    export_docx(blocks, docx_path, title=f"Протокол: {input_path.stem}")
    export_txt(corrected_text, txt_path)
    export_srt(blocks, srt_path)

    shutil.rmtree(workdir, ignore_errors=True)

    _log(progress_cb, "Готово.")
    return {"docx": docx_path, "txt": txt_path, "srt": srt_path}


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Транскрибация аудио/видео (RU/KZ/EN)")
    parser.add_argument("input", type=Path, help="Путь к аудио (mp3/wav/m4a) или видео (mp4) файлу")
    parser.add_argument("--output-dir", type=Path, default=Path("output"), help="Каталог для результатов")
    parser.add_argument("--diarize", action="store_true", help="Включить диаризацию спикеров (нужен HF_TOKEN)")
    parser.add_argument("--glossary", type=Path, default=Path("glossary.json"), help="Путь к glossary.json")
    parser.add_argument("--skip-llm", action="store_true", help="Пропустить LLM-коррекцию (для отладки ASR)")
    args = parser.parse_args()

    result = run_pipeline(
        args.input,
        args.output_dir,
        diarize_flag=args.diarize,
        glossary_path=args.glossary,
        skip_llm=args.skip_llm,
        progress_cb=print,
    )
    for kind, path in result.items():
        print(f"{kind}: {path}")


if __name__ == "__main__":
    main()
